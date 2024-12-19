from log import authLog
from docx import Document
from docx.shared import RGBColor
from auth import Auth
from commandsCLI import shCoreInfo, shIntDesSDW
from openpyxl import Workbook

import re
import os
import csv
import json
import traceback
import ipaddress
import openpyxl

removeCIDR_Patt = r'/\d{2}'
filterSiteCode = r'-sdw-0[1-9]'

PID_SDW03 = 'C8300-1N1S-4T2X-'
PID_SDW04 = 'C8300-1N1S-4T2X-'

ndlmPath1 = "NDLM_Template.xlsx"
ndlmPath2 = "NDLM_Tier4_Template.xlsx"

outputFolder = "Outputs"

sdw03Template = "Tier 4 - sdw-03 Template.csv"
sdw04Template = "Tier 4 - sdw-04 Template.csv"

returnList = []

def chooseCSV():
    csvDataList = []

    for i in range(2):
        while True:
            csvFile = input(f"Please enter the name of the CSV file for SDW-0{i + 1}: ")
            try:
                with open(csvFile, "r") as csvFileFinal:
                    authLog.info(f"User chose  the CSV File path: {csvFile}")
                    print(f"INFO: file successfully found.")
                    csvReader = csv.reader(csvFileFinal)
                    csvData = list(csvReader)
                    if csvData:
                        rowText = csvData[1]
                        for row in rowText:
                            print(f"{row}")
                        csvDataList.append(rowText)                         
                        break
                    else:
                        print(f"INFO: Cells not found under file: {csvFile}")
                        authLog.info(f"Cells not found under file: {csvFile}")
            except FileNotFoundError:
                print("File not found. Please check the file path and try again.")
                authLog.error(f"File not found in path {csvFile}")
                authLog.error(traceback.format_exc())
                continue

            except Exception as error:
                print(f"ERROR: {error}\n", traceback.format_exc())
                authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())
                
    mergedData = [item for sublist in csvDataList for item in sublist]
    # for index, item in enumerate(mergedData):
    #     print(f"rowText[{index}] with string: {item}")
    # os.system("PAUSE")
    return mergedData

def chooseDocx_vEdge(rowText):
    swHostname, username, netDevice = Auth(rowText[13])
    shHostnameOut, netVlan1105, netVlan1107, shIntDesSDWOut, shIntDesCONOut1, shVlanMgmtIP, shVlanMgmtCIDR, shLoop0Out = shCoreInfo(swHostname, username, netDevice)

    print(f"\n","="*76)
    print(f"INFO: Location: {rowText[3]}\n")

    print(f"INFO: BB1 Circuit Information: {rowText[25]}\n")

    print(f"="*76, "\n")

    while True:
        try:
            wordFile = "Tier IV - 8300 - vEdge to cEdge - 1 switch - gold.docx"
            wordDOC = Document(wordFile)
            authLog.info(f"User chose  the DOCX File path: {wordFile}")
            print(f"INFO: file successfully found: {wordFile}.")
            serialNumSDW03 = input("Please input the serial number of SDW-03: ")
            serialNumSDW04 = input("Please input the serial number of SDW-04: ")
            city = input("Please input the City: ")
            state = input("Please input the State: ")
            bb1Carrier = input("Please input the bb1-carrier: ")
            bb1Circuitid = input("Please input the bb1-circuitid: ")
            cEdge1TLOC3_Port = input(f"Please input the cedge1-tloc3-port (TenGigabitEthernet0/0/5 or GigabitEthernet0/0/1 for {bb1Carrier} port): ")
            print("=" * 61,"\n\tINFO: Now begins information of the Core Switch")
            print("=" * 61)
            print(f"{shHostnameOut}{shIntDesSDW}\n{shIntDesSDWOut}\n")
            swcEdge1_vlan = input("Please input the VLAN for SDW-03, 1105 if possible: ")
            swcEdge2_vlan = input("Please input the VLAN for SDW-04, 1107 if possible: ")
            swcEdge1_port = input("Please input the connection to SDW-03 gi0/0/0 in VPN 1 from the switch: ")
            swcEdge2_port = input("Please input the connection to SDW-04 gi0/0/0 in VPN 1 from the switch: ")
            
            # print("\nrowText 2:", rowText[2], "rowText 20:", rowText[20])
            # print("After changes:")
            # This one changes the hostname from 01-02 to 03-04
            # Example: nyny-cen-sdw-02 - ge0/3 - TLOC1 Ext to nyny-cen-sdw-04 - ge0/0/3 - TLOC1 Ext
            rowText[2] = re.sub('01', '03', rowText[2])
            rowText[18] = re.sub('02', '04', rowText[18])
            rowText[18] = re.sub('ge0/3', 'ge0/0/3', rowText[18])
            # print("rowText 2:", rowText[2], "rowText 20:", rowText[20])
            # os.system("PAUSE")

            # print("\nrowText 47:", rowText[47], "rowText 65:", rowText[65])
            # print("After changes:")
            # This one changes the hostname from 01-02 to 03-04
            # Example: nyny-cen-sdw-01 - ge0/3 - TLOC1 to nyny-cen-sdw-03 - ge0/0/3 - TLOC1
            rowText[43] = re.sub('02', '04', rowText[43])
            rowText[60] = re.sub('01', '03', rowText[60])
            rowText[60] = re.sub('ge0/3', 'ge0/0/3', rowText[60])
            # print("rowText 47:", rowText[47], "rowText 65:", rowText[65])
            # os.system("PAUSE")

            # print("rowText 6:", rowText[6], "rowText 21:", rowText[21], "rowText 32:", rowText[32], \
            #       f"rowText 51:", rowText[51], "rowText 85:", rowText[85],"")
            # print("After changes:")
            # Removes /30
            rowText[6] = re.sub(removeCIDR_Patt, '', rowText[6]) # This is for cedge1-rtr-ip
            rowText[19] = re.sub(removeCIDR_Patt, '', rowText[19])
            rowText[47] = re.sub(removeCIDR_Patt, '', rowText[47]) # This is for cedge2-rtr-ip
            rowText[61] = re.sub(removeCIDR_Patt, '', rowText[61])
            # print("rowText 6:", rowText[6], "rowText 21:", rowText[21], "rowText 32:", rowText[32], \
            #       f"rowText 51:", rowText[51], "rowText 85:", rowText[85],"")
            # os.system("PAUSE")

            cedge1TLOC3_List = rowText[26] # cedge1-tloc3-ip/cedge1-tloc3-cidr
            cedge1TLOC3_STR = ''.join(cedge1TLOC3_List)
            cedge1TLOC3_IP_STR = cedge1TLOC3_STR.split('/')[0]
            cedge1TLOC3_CIDR_STR = cedge1TLOC3_STR.split('/')[1]
            cedge1TLOC3_MASK_STR = ipaddress.IPv4Network(cedge1TLOC3_STR, strict=False)
            cedge1TLOC3_MASK_STR = str(cedge1TLOC3_MASK_STR.netmask)

            serialNumSDW03New = PID_SDW03 + serialNumSDW03
            serialNumSDW04New = PID_SDW04 + serialNumSDW04

            siteCode = f'{rowText[2]}'
            siteCode = re.sub(filterSiteCode, '', siteCode)
            print(f"This is the side code:{siteCode}")
            # os.system("PAUSE")
            sw_host = f'{rowText[13]}'
            serialNumSDW01 = f'{rowText[0]}'
            serialNumSDW02 = f'{rowText[41]}'
            cEdge1Loop = f'{rowText[1]}'
            cEdge2Loop = f'{rowText[42]}'
            siteNo = f'{rowText[40]}'

            replaceText = {
                'cedge1-loop' : f'{rowText[1]}',  # OK
                'cedge1-host' : f'{rowText[2]}',  # OK
                'snmp-location' : f'{rowText[3]}', # OK
                'cedge1-rtr-ip' : f'{rowText[6]}', # OK
                'cEdge-asn' : f'{rowText[9]}', # OK
                'cedge1-sw-ip' : f'{rowText[12]}', # OK
                'switch-asn' : f'{rowText[14]}', # OK
                'cedge1-tloc3-gate' : f'{rowText[16]}', # OK
                'cedge1-tloc3-ext-ip' : f'{rowText[19]}', # OK
                'bb1-up-speed' : f'{rowText[34]}', # OK
                'bb1-down-speed' : f'{rowText[35]}', # OK
                'latitude' : f'{rowText[37]}', # OK
                'longitude' : f'{rowText[38]}', # OK
                'site-no'	: f'{rowText[40]}',
                # Here starts the second CSV file #
                'cedge2-loop' : f'{rowText[42]}', # OK
                'cedge2-host'	: f'{rowText[43]}', # OK
                'cedge2-rtr-ip' : f'{rowText[47]}', # OK
                'cedge2-sw-ip' : f'{rowText[53]}', # OK
                'cedge2-tloc3-ip' : f'{rowText[61]}', # OK
                'cellular-up-speed' : f'{rowText[67]}', # OK
                'cellular-down-speed' : f'{rowText[68]}', # OK
            }

            # print(json.dumps(replaceText, indent=4))
            # os.system("PAUSE")

            stringRegexPatt = {
                'city': city, # OK
                'state': state, # OK
                'site-code': siteCode, #nyny or aztuc-lan
                'sw-mgmt-ip' : shVlanMgmtIP, # OK
                'bb1-carrier': bb1Carrier, # OK
                'bb1-circuitid': bb1Circuitid, # OK
                'cedge1-tloc3-port': cEdge1TLOC3_Port, # OK
                'cedge1-tloc3-ip': cedge1TLOC3_IP_STR, # OK
                'cedge1-tloc3-mask' : cedge1TLOC3_MASK_STR, # OK
                'cedge1-tloc3-cidr': cedge1TLOC3_CIDR_STR, # OK
                'cedge1-lan-net': netVlan1105, # OK
                'cedge2-lan-net': netVlan1107, # OK
                'sw-loop': shLoop0Out, # OK
                'sw-mgmt-cidr': shVlanMgmtCIDR, # OK
                'sw-cedge1-port': swcEdge1_port,  # OK
                'sw-cedge1-vlan': swcEdge1_vlan, # OK
                'sw-cedge2-port': swcEdge2_port, # OK
                'sw-cedge2-vlan': swcEdge2_vlan, # OK
                'sw-remote-con-net1': shIntDesCONOut1[0], # OK
                'sw-remote-con-net2': shIntDesCONOut1[1], # OK
                'sw-host' : sw_host, # OK
                'sw-mgmt-vlan' : '1500', # OK
                'cedge1-serial-no' : serialNumSDW03New, # OK
                'cedge2-serial-no' : serialNumSDW04New # OK
            }

            manualReplacements = {re.compile(r'\b{}\b'.format(pattern), re.IGNORECASE): value for pattern, value in stringRegexPatt.items()}

            for para in wordDOC.paragraphs:
                if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in para.runs):
                    print(f"Found red text: {para.text}")
                    for wordString, csvString in replaceText.items():
                        if re.search(r'\b{}\b'.format(re.escape(wordString)), para.text, re.IGNORECASE):
                            print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                            authLog.info(f"Replacing '{wordString}' with '{csvString}'")
                            para.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, para.text, flags=re.IGNORECASE)

                    for placeholder, replacement in manualReplacements.items():
                        replacement = str(replacement)
                        if placeholder.search(para.text):
                            print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            authLog.info(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            para.text = placeholder.sub(replacement, para.text)

            for table in wordDOC.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in paragraph.runs):
                                print(f"Found red text: {paragraph.text}")
                                for wordString, csvString in replaceText.items():
                                    if re.search(r'\b{}\b'.format(re.escape(wordString)), paragraph.text, re.IGNORECASE):
                                        print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                                        authLog.info(f"Replacing in Table: '{wordString}' with '{csvString}'")
                                        paragraph.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, paragraph.text, flags=re.IGNORECASE)

                                for placeholder, replacement in manualReplacements.items():
                                    replacement = str(replacement)
                                    if placeholder.search(paragraph.text):
                                        print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                                        authLog.info(f"Replacing in Table: '{placeholder.pattern}' with '{replacement}'")
                                        paragraph.text = placeholder.sub(replacement, paragraph.text)

            newWordDoc = f"Outputs/{siteCode} - vEdge to cEdge Implementation Plan.docx"
            wordDOC.save(newWordDoc)
            authLog.info(f"Replacements made successfully in DOCX file and saved as: {newWordDoc}")
            print(f"INFO: Replacements made successfully in DOCX file and saved as: {newWordDoc}")
            
            os.system("PAUSE")

            ignored = "EMPTY"

            manualReplaceList = [
                serialNumSDW01,     #0
                serialNumSDW02,     #1
                serialNumSDW03New,  #2
                serialNumSDW04New,  #3
                cEdge1Loop,         #4
                cEdge2Loop,         #5
                siteNo,             #6
                city,               #7
                state,              #8
                siteCode,           #9
                shVlanMgmtIP,       #10            
                ignored,            #11
                ignored,            #12
                ignored,            #13
                bb1Carrier,         #14
                bb1Circuitid,       #15
                cEdge1TLOC3_Port,   #16
                cedge1TLOC3_IP_STR, #17
                cedge1TLOC3_MASK_STR,#18
                cedge1TLOC3_CIDR_STR,#19
                netVlan1105,        #20
                netVlan1107,        #21
                shLoop0Out,         #22
                shVlanMgmtCIDR,     #23
                swcEdge1_port,      #24
                swcEdge1_vlan,      #25
                swcEdge2_port,      #26
                swcEdge2_vlan,      #27
                ignored,            #28
                shIntDesCONOut1[0], #29
                shIntDesCONOut1[1], #30
                sw_host,            #31
                '1500'              #32

            ]

            return {
                'rowText' : rowText,
                'rowText1' :  manualReplaceList
            }

        
        except FileNotFoundError:
            print("File not found. Please check the file path and try again.")
            authLog.error(f"File not found in path {wordFile}")
            authLog.error(traceback.format_exc())
            continue

        except Exception as error:
            print(f"ERROR: {error}\n", traceback.format_exc())
            authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def modNDLMvEdge(rowText, rowText1):
    try:
        cedge1_serial_no = f'{rowText1[2]}'
        cedge1_serial_no = re.sub(PID_SDW03, '', cedge1_serial_no)
        cedge2_serial_no = f'{rowText1[3]}'
        cedge2_serial_no = re.sub(PID_SDW04, '', cedge2_serial_no)
        replaceText = {
            'site-code' : f'{rowText1[9]}', # OK
            'vedge1-serial-no' : f'{rowText1[0]}', # OK
            'vedge2-serial-no' : f'{rowText1[1]}', # OK
            'cedge1-serial-no' : cedge1_serial_no, # OK
            'cedge2-serial-no' : cedge2_serial_no, # OK
            'cedge1-loop' : f'{rowText1[4]}', # OK
            'cedge2-loop' : f'{rowText1[5]}', # OK
            'snmp-location' : f'{rowText[3]}', # OK
            'vedge1-loop': f'{rowText1[4]}', # OK 
            'vedge2-loop': f'{rowText1[5]}' # OK
        }

        ndlmFile = openpyxl.load_workbook(ndlmPath1)
        ndlmFileSheet = ndlmFile.active

        for row in ndlmFileSheet.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                            authLog.info(f"Replacing '{key}' with '{value}' in the NDLM File 1")
                    cell.value = cellValue

            newNDLMFile = os.path.join(outputFolder, f'{rowText1[9]} - NDLM.xlsx')
            ndlmFile.save(newNDLMFile)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath1}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def modNDLM2vEdge(rowText, rowText1):
    # print(f'mpls-speed: {rowText[38]}')
    # print(f'bb1-up-speed : {rowText[82]}')
    # print(f'bb1-down-speed : {rowText[82]}')
    # print(f'bb1-carrier: {rowText1[14]}')
    # os.system("PAUSE")

    try:

        replaceText = {
            'site-code' : f'{rowText1[9]}', # OK
            'cedge1-loop' : f'{rowText1[4]}', # OK
            'cedge2-loop' : f'{rowText1[5]}', # OK
            'snmp-location' : f'{rowText[3]}', # OK
            'site-no': f'{rowText1[6]}', # OK
            'city': f'{rowText1[7]}', # OK
            'state': f'{rowText1[8]}', # OK
            'cedge1-host': f'{rowText[2]}', # OK
            'cedge2-host': f'{rowText[43]}', # OK
            'sw-host' : f'{rowText1[31]}', # OK
            'cedge1-tloc3-port': f'{rowText1[16]}', # OK
            'sw-cedge1-port' : f'{rowText1[24]}', # OK
            'sw-cedge2-port' : f'{rowText1[26]}', # OK
            'bb1-up-speed' : f'{rowText[34]}',  # OK
            'bb1-down-speed' : f'{rowText[35]}',  # OK
            'bb1-carrier' : f'{rowText1[14]}'  # OK
        }

        ndlmFile1 = openpyxl.load_workbook(ndlmPath2)
        ndlmFileSheet1 = ndlmFile1.active

        for row in ndlmFileSheet1.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                            authLog.info(f"Replacing '{key}' with '{value}' in the NDLM File 2")
                    cell.value = cellValue

            newNDLMFile1 = os.path.join(outputFolder, f'{rowText1[9]} - NDLM - Tier4.xlsx')
            ndlmFile1.save(newNDLMFile1)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath2}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def cEdgeTemplatevEdge(rowText, rowText1):

    # for index, item in enumerate(rowText):
    #     print(f"rowText[{index}] with string: {item}")
    
    # for index, item in enumerate(rowText1):
    #     print(f"rowText1[{index}] with string: {item}")
    # os.system("PAUSE")
    
    newSDW03Template = f'Outputs/{rowText1[9]} - SDW-03 Template.csv'
    newSDW04Template = f'Outputs/{rowText1[9]} - SDW-04 Template.csv'

    templateReplacements = {
        'cedge1-loop' : f'{rowText[1]}',  # OK
        'cedge1-host' : f'{rowText[2]}',  # OK
        'snmp-location' : f'{rowText[3]}', # OK
        'cedge1-rtr-ip' : f'{rowText[6]}', # OK
        'cEdge-asn' : f'{rowText[9]}', # OK
        'cedge1-sw-ip' : f'{rowText[12]}', # OK
        'switch-asn' : f'{rowText[14]}', # OK
        'cedge1-tloc3-gate' : f'{rowText[16]}', # OK
        'cedge1-tloc3-ext-ip' : f'{rowText[19]}', # OK
        'bb1-up-speed' : f'{rowText[34]}', # OK
        'bb1-down-speed' : f'{rowText[35]}', # OK
        'latitude' : f'{rowText[37]}', # OK
        'longitude' : f'{rowText[38]}', # OK
        'site-no'	: f'{rowText[40]}',
        # Here starts the second CSV file #
        'cedge2-loop' : f'{rowText[42]}', # OK
        'cedge2-host'	: f'{rowText[43]}', # OK
        'cedge2-rtr-ip' : f'{rowText[47]}', # OK
        'cedge2-sw-ip' : f'{rowText[53]}', # OK
        'cedge2-tloc3-ip' : f'{rowText[61]}', # OK
        'cellular-up-speed' : f'{rowText[67]}', # OK
        'cellular-down-speed' : f'{rowText[68]}', # OK

        'cedge1-serial-no' : rowText1[2], # OK
        'cedge2-serial-no' : rowText1[3], # OK
        'cedge1-loop' : rowText1[4], # OK
        'cedge2-loop' : rowText1[5], # OK
        'site-no'	: rowText1[6], # OK
        'city': rowText1[7], # OK 
        'state': rowText1[8], # OK
        'site-code': rowText1[9], # OK
        'sw-mgmt-ip' : rowText1[10], # OK
        'bb1-carrier': rowText1[14], # OK
        'bb1-circuitid': rowText1[15], # OK
        'cedge1-tloc3-port': rowText1[16], # OK
        'cedge1-tloc3-ip': rowText1[17], # OK
        'cedge1-tloc3-mask' : rowText1[18], # OK
        'cedge1-tloc3-cidr': rowText1[19], # OK
        'cedge1-lan-net': rowText1[20], # OK
        'cedge2-lan-net': rowText1[21], # OK
        'sw-loop': rowText1[22], # OK
        'sw-mgmt-cidr': rowText1[23], # OK
        'sw-cedge1-port': rowText1[24], # OK
        'sw-cedge1-vlan': rowText1[25], # OK
        'sw-cedge2-port': rowText1[26], # OK
        'sw-cedge2-vlan': rowText1[27], # OK
        'sw-remote-con-net1': rowText1[29], # OK
        'sw-remote-con-net2': rowText1[30], # OK
        'sw-host' : rowText1[31], # OK
        'sw-mgmt-vlan' : rowText1[32] # OK
    }

    try:
        with open(sdw03Template, "r") as inputCSV:
            authLog.info(f"Generating {rowText1[9]}-SDW-03-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-03-Template.")
            csvReader = csv.reader(inputCSV)
               
            rows = list(csvReader)

            if len(rows) > 1:
                secondRow = rows[1]
                modifiedRow = []
                for index, cell in enumerate(secondRow):
                    cellValue = str(cell).strip()
                    originalCellValue = cellValue
                    for key, value in templateReplacements.items():
                        if key.lower() in cellValue.lower():
                            cellValue = re.sub(re.escape(key), value, cellValue, flags=re.IGNORECASE)
                            authLog.info(f"Replacing '{key}' with '{value}' in row 2, cell {index + 1}:" \
                                         f"'{originalCellValue}' -> '{cellValue}', in the SDW-03-Template")
                    modifiedRow.append(cellValue)
                rows[1] = modifiedRow

        with open(newSDW03Template, 'w', newline="") as outputCSV:
            csvWriter = csv.writer(outputCSV)
            csvWriter.writerows(rows)
    
        with open(sdw04Template, "r") as inputCSV1:
            authLog.info(f"Generating {rowText1[9]}-SDW-04-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-04-Template.")
            csvReader1 = csv.reader(inputCSV1)
               
            rows1 = list(csvReader1)

            if len(rows1) > 1:
                secondRow1 = rows1[1]
                modifiedRow1 = []
                for index1, cell1 in enumerate(secondRow1):
                    cellValue1 = str(cell1).strip()
                    originalCellValue1 = cellValue1

                    for key1, value1 in templateReplacements.items():
                        if key1.lower() in cellValue1.lower():
                            cellValue1 = re.sub(re.escape(key1), value1, cellValue1, flags=re.IGNORECASE)
                            authLog.info(f"Replacing '{key1}' with '{value1}' in row 2, cell {index1 + 1}:" \
                                         f"'{originalCellValue1}' -> '{cellValue1}', in the SDW-04-Template")
                    modifiedRow1.append(cellValue1)
                rows1[1] = modifiedRow1
            
        with open(newSDW04Template, 'w', newline="") as outputCSV1:
            csvWriter1 = csv.writer(outputCSV1)
            csvWriter1.writerows(rows1)

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Error message: {error}\n", traceback.format_exc())